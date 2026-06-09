import os
import json
from datetime import date
from database.db_manager import DBManager
from core.knowledge_manager import KnowledgeManager
from utils.exporter import export_markdown, export_html, export_pdf

class CVTailor:
    def __init__(self, db_manager: DBManager, km: KnowledgeManager, llm_provider, config: dict):
        self.db = db_manager
        self.km = km
        self.llm = llm_provider
        self.config = config

    def build_cv_context(self, job: dict, retrieved_chunks: list) -> str:
        context = f"Job Title: {job.get('title', '')}\n"
        context += f"Job Description: {job.get('description', '')}\n\n"
        context += "Evidence from Candidate Knowledge Base:\n"
        for i, chunk in enumerate(retrieved_chunks):
            context += f"[{i+1}] Source: {chunk['metadata'].get('source', 'Unknown')} - {chunk['text']}\n"
        return context

    def generate_profile_and_skills(self, job: dict, evidence_chunks: list) -> dict:
        context = self.build_cv_context(job, evidence_chunks)
        prompt = f"""
You are an expert resume writer. Generate ONLY the 'Profile' and 'Skills' sections of a CV tailored to this job.
Rules:
1. Base your writing strictly on the provided Evidence. Do not hallucinate.
2. Provide your output as JSON with keys "profile" (string) and "skills" (list of strings).
3. If evidence is missing for a requirement, write a weaker but honest sentence, or exclude it.

Context:
{context}
"""
        response = self.llm.generate_completion(prompt)
        try:
            import re
            match = re.search(r'\{.*\}', response, re.DOTALL)
            if match:
                res = json.loads(match.group(0))
            else:
                res = json.loads(response)
                
            return {
                "profile": res.get("profile", "Professional candidate tailored to the role."),
                "skills": res.get("skills", ["Relevant Skill"])
            }
        except Exception:
            return {"profile": "Experienced professional tailored to the role.", "skills": ["Relevant Skill"]}

    def assemble_cv(self, profile: str, skills: list, static_experience: str, static_education: str, evidence_sources: list) -> str:
        cv_md = f"# Tailored CV\n\n"
        cv_md += f"## Profile\n{profile}\n\n"
        cv_md += f"## Skills\n"
        for skill in skills:
            cv_md += f"- {skill}\n"
        cv_md += f"\n## Experience\n{static_experience}\n\n"
        cv_md += f"## Education\n{static_education}\n\n"
        
        cv_md += f"## Evidence Sources\n"
        for src in evidence_sources:
            cv_md += f"- {src}\n"
            
        return cv_md

    def generate_tailored_cv(self, job_id: str):
        jobs = self.db.execute_query("SELECT * FROM jobs WHERE job_id = ?", (job_id,))
        if not jobs:
            print(f"Job {job_id} not found.")
            return False
            
        job = jobs[0]
        if job['status'] not in ['needs_review', 'approved']:
            print(f"Job {job_id} is in status {job['status']}, not generating CV.")
            return False

        # Read user settings
        settings_path = "knowledge_base/settings.json"
        cv_rules = []
        base_cv_path = "profile/base_cv.md"
        
        if os.path.exists(settings_path):
            try:
                with open(settings_path, 'r', encoding='utf-8') as f:
                    st = json.load(f)
                    cv_rules = st.get("cv_rules", [])
                    if st.get("default_cv_template"):
                        base_cv_path = os.path.join("knowledge_base", "processed_sources", st["default_cv_template"])
            except Exception:
                pass
                
        if not os.path.exists(base_cv_path):
            print(f"Base CV not found at {base_cv_path}")
            # Try fallback
            fallback = os.path.join("knowledge_base", "processed_sources", "base_cv_986a0271.md")
            if os.path.exists(fallback):
                base_cv_path = fallback
            else:
                return False
                
        if base_cv_path.endswith('.pdf'):
            import pypdf
            with open(base_cv_path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                base_cv_content = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
        else:
            with open(base_cv_path, 'r', encoding='utf-8', errors='replace') as f:
                base_cv_content = f.read()

        # Query RAG
        query = f"Job title: {job.get('title', '')}. Description: {job.get('description', '')}"
        rag_results = self.km.query_knowledge_base(query, top_k=5)
        
        context = ""
        for i, chunk in enumerate(rag_results):
            context += f"[{i+1}] Source: {chunk['metadata'].get('source', 'Unknown')} - {chunk['text']}\n"
            
        rules_text = "\n".join([f"- {r}" for r in cv_rules]) if cv_rules else "- No special formatting rules provided."
        
        prompt = f"""
You are an expert resume writer.

### Job Description:
```
{job.get('description', '')}
```

### Evidence from Candidate Knowledge Base:
{context}

### Strict Custom Rules from User:
{rules_text}

### INSTRUCTIONS:
You need to generate two highly tailored sections for the candidate's CV based on the Job Description and Evidence.
1. "profile": A short, impactful professional summary (3-4 sentences).
2. "skills": A bulleted list of skills relevant to the job, formatted as markdown bullets.

You MUST STRICTLY FOLLOW the "Strict Custom Rules from User". Do not hallucinate skills.

Output ONLY a valid JSON object with the keys "profile" and "skills".
"""
        response = self.llm.generate_completion(prompt)
        
        import json
        import re
        
        # Clean JSON markdown block if exists
        response_clean = re.sub(r'^```(?:json)?\s*', '', response.strip())
        response_clean = re.sub(r'\s*```$', '', response_clean)
        
        try:
            cv_data = json.loads(response_clean)
        except Exception as e:
            print(f"Failed to parse JSON: {e}")
            cv_data = {"profile": "Failed to generate profile.", "skills": "Failed to generate skills."}
            
        # Outputs directory structure
        from datetime import date
        today = date.today().strftime("%Y-%m-%d")
        safe_company = "".join(x for x in (job.get('company') or "Unknown") if x.isalnum() or x in " _-")
        safe_title = "".join(x for x in (job.get('title') or "Job") if x.isalnum() or x in " _-")
        folder_name = f"{safe_company}_{safe_title}".replace(" ", "_")
        if not folder_name:
            folder_name = "Unknown_Job"
            
        out_dir = os.path.join("outputs", today, folder_name)
        os.makedirs(out_dir, exist_ok=True)
        
        # Ensure skills and profile are strings
        skills_val = cv_data.get('skills', '')
        if isinstance(skills_val, list):
            skills_val = '\n'.join([f"- {item}" for item in skills_val])
        else:
            skills_val = str(skills_val)
            
        profile_val = cv_data.get('profile', '')
        if isinstance(profile_val, list):
            profile_val = ' '.join([str(item) for item in profile_val])
        else:
            profile_val = str(profile_val)
            
        is_docx = base_cv_path.endswith('.docx')
        
        if is_docx:
            import docx
            from docx2pdf import convert as docx2pdf_convert
            doc = docx.Document(base_cv_path)
            
            # Replace placeholders
            for p in doc.paragraphs:
                if '{{PROFILE}}' in p.text:
                    p.text = p.text.replace('{{PROFILE}}', profile_val)
                if '{{SKILLS}}' in p.text:
                    # Skills usually have newlines, docx paragraphs don't handle newlines well directly in text replace without splitting
                    p.text = p.text.replace('{{SKILLS}}', skills_val.replace('\n', '\n'))
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if '{{PROFILE}}' in p.text:
                                p.text = p.text.replace('{{PROFILE}}', profile_val)
                            if '{{SKILLS}}' in p.text:
                                p.text = p.text.replace('{{SKILLS}}', skills_val)
                                
            final_path = os.path.join(out_dir, "tailored_cv.docx")
            doc.save(final_path)
            
            pdf_path = os.path.join(out_dir, "tailored_cv.pdf")
            try:
                # Requires MS Word installed on Windows
                docx2pdf_convert(final_path, pdf_path)
            except Exception as e:
                print(f"Failed to convert DOCX to PDF: {e}")
        else:
            # Fallback to saving markdown with the replaced text
            cv_md = base_cv_content.replace('{{PROFILE}}', profile_val).replace('{{SKILLS}}', skills_val)
            final_path = os.path.join(out_dir, "tailored_cv.md")
            export_markdown(cv_md, final_path)
            html_path = os.path.join(out_dir, "tailored_cv.html")
            pdf_path = os.path.join(out_dir, "tailored_cv.pdf")
            export_html(final_path, html_path)
            export_pdf(html_path, pdf_path)
            
        # Update DB
        self.db.execute_query(
            "UPDATE jobs SET status = 'cv_generated', generated_cv_path = ? WHERE job_id = ?",
            (final_path, job_id)
        )
        
        # Add to knowledge manager for future reuse (we just save the JSON as text representation)
        self.km.add_generated_cv(json.dumps(cv_data), job_id)
        
        return True
