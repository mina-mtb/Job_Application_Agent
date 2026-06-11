import docx
import logging

def find_and_replace_in_runs(paragraph, replacements, found_placeholders):
    """
    Search for keys in `replacements` within the paragraph.
    If found, replace it while attempting to preserve formatting by 
    replacing within a single run if possible, or falling back to 
    replacing the whole paragraph text.
    """
    for search_text, replace_text in replacements.items():
        if search_text in paragraph.text:
            found_placeholders.add(search_text)
            
            # Fast/clean path: if the search_text is entirely within a single run
            replaced_in_run = False
            for run in paragraph.runs:
                if search_text in run.text:
                    # Fix newlines for docx
                    run.text = run.text.replace(search_text, replace_text)
                    replaced_in_run = True
                    break
            
            # Fallback: if search_text spans multiple runs, fallback to rewriting the paragraph text
            if not replaced_in_run:
                paragraph.text = paragraph.text.replace(search_text, replace_text)
                
def render_docx_template(template_path: str, output_path: str, replacements: dict) -> dict:
    """
    Replaces dictionary keys in `replacements` with their values in the given DOCX template.
    Searches paragraphs, tables, headers, and footers.
    
    Returns:
    {
        "success": bool,
        "found_placeholders": list,
        "missing_placeholders": list,
        "output_path": str,
        "warnings": list
    }
    """
    result = {
        "success": False,
        "found_placeholders": [],
        "missing_placeholders": [],
        "output_path": output_path,
        "warnings": []
    }
    
    try:
        doc = docx.Document(template_path)
    except Exception as e:
        result["warnings"].append(f"Failed to load document: {e}")
        return result
        
    found = set()
    
    # 1. Paragraphs
    for p in doc.paragraphs:
        find_and_replace_in_runs(p, replacements, found)
        
    # 2. Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    find_and_replace_in_runs(p, replacements, found)
                    
    # 3. Headers and Footers
    for section in doc.sections:
        for p in section.header.paragraphs:
            find_and_replace_in_runs(p, replacements, found)
        for p in section.footer.paragraphs:
            find_and_replace_in_runs(p, replacements, found)
            
    result["found_placeholders"] = list(found)
    
    for req_placeholder in replacements.keys():
        if req_placeholder not in found:
            result["missing_placeholders"].append(req_placeholder)
            
    if result["missing_placeholders"]:
        result["warnings"].append(f"Missing required placeholders: {', '.join(result['missing_placeholders'])}")
        # Fail clearly if required placeholder missing (as requested)
        return result
        
    try:
        doc.save(output_path)
        result["success"] = True
    except Exception as e:
        result["warnings"].append(f"Failed to save document: {e}")
        
    return result
