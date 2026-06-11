import os
import docx
import pytest
from core.docx_template_renderer import render_docx_template

def create_test_template(path):
    doc = docx.Document()
    doc.add_paragraph('Hello {{PROFILE}}')
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = 'Here are my {{SKILLS}}'
    doc.save(path)

def test_docx_template_renderer(tmp_path):
    template_path = str(tmp_path / "template.docx")
    output_path = str(tmp_path / "output.docx")
    
    create_test_template(template_path)
    
    replacements = {
        "{{PROFILE}}": "John Doe",
        "{{SKILLS}}": "Python, pytest"
    }
    
    result = render_docx_template(template_path, output_path, replacements)
    
    assert result["success"] is True
    assert "{{PROFILE}}" in result["found_placeholders"]
    assert "{{SKILLS}}" in result["found_placeholders"]
    assert len(result["missing_placeholders"]) == 0
    assert os.path.exists(output_path)
    
    # Check output content
    out_doc = docx.Document(output_path)
    text = "\n".join(p.text for p in out_doc.paragraphs)
    for table in out_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += "\n".join(p.text for p in cell.paragraphs)
                
    assert "John Doe" in text
    assert "Python, pytest" in text
    assert "{{PROFILE}}" not in text
    assert "{{SKILLS}}" not in text
    
def test_missing_placeholder(tmp_path):
    template_path = str(tmp_path / "template_missing.docx")
    output_path = str(tmp_path / "output_missing.docx")
    
    doc = docx.Document()
    doc.add_paragraph('Hello {{PROFILE}}')
    doc.save(template_path)
    
    replacements = {
        "{{PROFILE}}": "John Doe",
        "{{SKILLS}}": "Python, pytest"
    }
    
    result = render_docx_template(template_path, output_path, replacements)
    
    # Should fail because {{SKILLS}} is missing
    assert result["success"] is False
    assert "{{PROFILE}}" in result["found_placeholders"]
    assert "{{SKILLS}}" in result["missing_placeholders"]
    assert len(result["warnings"]) > 0
