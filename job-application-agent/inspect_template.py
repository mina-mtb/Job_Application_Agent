import docx
import sys
sys.stdout.reconfigure(encoding='utf-8')

doc = docx.Document('knowledge_base/processed_sources/Mina_Tahmasebi_Skills_Profile.docx')
print('=== PARAGRAPHS ===')
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f'P{i}: [{p.text[:100]}]')
    if '{{' in p.text:
        print(f'  *** PLACEHOLDER FOUND: {p.text}')

print('\n=== TABLES ===')
for ti, table in enumerate(doc.tables):
    print(f'Table {ti}: {len(table.rows)} rows x {len(table.columns)} cols')
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            txt = cell.text.strip()
            if txt:
                print(f'  [{ri},{ci}]: {txt[:100]}')
            if '{{' in cell.text:
                print(f'  *** PLACEHOLDER FOUND in cell [{ri},{ci}]: {cell.text}')

print('\n=== ALSO CHECK TEMPLATE FILE ===')
doc2 = docx.Document('knowledge_base/processed_sources/Mina TAhmasebi Cv Template (1).docx')
print('--- Paragraphs ---')
for i, p in enumerate(doc2.paragraphs):
    if p.text.strip():
        print(f'P{i}: [{p.text[:100]}]')
    if '{{' in p.text:
        print(f'  *** PLACEHOLDER FOUND: {p.text}')

print('--- Tables ---')
for ti, table in enumerate(doc2.tables):
    print(f'Table {ti}: {len(table.rows)} rows x {len(table.columns)} cols')
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            txt = cell.text.strip()
            if txt:
                print(f'  [{ri},{ci}]: {txt[:100]}')
            if '{{' in cell.text:
                print(f'  *** PLACEHOLDER FOUND in cell [{ri},{ci}]: {cell.text}')
