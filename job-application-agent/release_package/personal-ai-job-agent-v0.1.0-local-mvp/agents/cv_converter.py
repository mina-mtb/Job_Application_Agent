import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches

def markdown_to_docx(md_path, docx_path):
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
        
    # Strip frontmatter if present
    if content.startswith('---'):
        parts = content.split('---', 2)
        if len(parts) >= 3:
            content = parts[2].strip()

    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_runs(p, line[2:])
        elif line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_runs(p, line[2:])
        else:
            p = doc.add_paragraph()
            _add_formatted_runs(p, line)
            
    doc.save(docx_path)

def _add_formatted_runs(paragraph, text):
    # Extract **bold**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Extract *italic* or _italic_
            subparts = re.split(r'(_.*?_|\*.*?\*)', part)
            for sub in subparts:
                if (sub.startswith('_') and sub.endswith('_')) or (sub.startswith('*') and sub.endswith('*')):
                    if len(sub) > 2:
                        run = paragraph.add_run(sub[1:-1])
                        run.italic = True
                else:
                    paragraph.add_run(sub)

def main():
    target_dir = Path(__file__).parent.parent / "cvs" / "tailored"
    if not target_dir.exists():
        print(f"Directory {target_dir} not found.")
        return
        
    md_files = list(target_dir.glob("*.md"))
    print(f"Found {len(md_files)} Markdown CVs.")
    
    for md_file in md_files:
        docx_file = md_file.with_suffix('.docx')
        print(f"Converting {md_file.name} -> {docx_file.name}")
        markdown_to_docx(md_file, docx_file)
        
    print("Done!")
        
if __name__ == "__main__":
    main()
